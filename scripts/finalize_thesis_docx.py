from docx import Document


INPUT_PATH = r"D:\CSEdge_Internship_Thesis_Satyam_Kumar.docx"
OUTPUT_PATH = r"D:\CSEdge_Internship_Thesis_Satyam_Kumar_Final.docx"


def add_paragraphs(doc, items, style="Normal"):
    for text in items:
        p = doc.add_paragraph(text)
        p.style = style


def main():
    doc = Document(INPUT_PATH)

    # Fill remaining known placeholders without forcing unknown values.
    placeholder_map = {
        "[Your Faculty Guide Name]": "Prof. [Internal Guide Name]",
        "[Your Internship Mentor Name]": "[Industry Guide Name]",
        "[Your Internship Company Name], [City], [State]": "[Internship Company], [City], [State]",
    }

    for p in doc.paragraphs:
        t = p.text
        if not t:
            continue
        updated = t
        for old, new in placeholder_map.items():
            updated = updated.replace(old, new)
        if updated != t:
            if p.runs:
                p.runs[0].text = updated
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = updated

    # Append expanded sections for submission readiness and page-length growth.
    h1 = "Heading 1"
    h2 = "Heading 2"

    doc.add_page_break()
    p = doc.add_paragraph("9. Expanded Implementation Details")
    p.style = h1

    p = doc.add_paragraph("9.1 Frontend Architecture and Routing")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "The frontend application is developed using React with Vite and follows a component-driven architecture for maintainability and scalability.",
            "Routing is managed through React Router with role-aware protected routes to ensure that student, educator, and admin workflows remain isolated and secure.",
            "Core pages include Home, Course List, Course Details, My Enrollments, Player, Aptitude/DSA/Dev/SQL practice pages, Arena page, and dashboard modules.",
            "Reusable UI components such as cards, hero sections, navigation, and feedback toasts improve consistency and reduce duplicate code.",
            "The frontend communicates with backend APIs using Axios instances with token-based authorization headers managed through Clerk authentication context.",
        ],
    )

    p = doc.add_paragraph("9.2 Backend Architecture and API Layer")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "The backend is built using Node.js and Express with modular route-controller-service organization.",
            "Each feature area (user, course, test, educator, admin, chat, sprint, resume, and study-share) is exposed through dedicated API route groups.",
            "Controllers process business logic such as course publication, enrollment validation, test scoring, and role-based access checks.",
            "Data persistence is handled through MongoDB using Mongoose schemas for structured model validation.",
            "Cross-origin configuration and robust middleware flow ensure secure and stable communication between frontend and backend services.",
        ],
    )

    p = doc.add_paragraph("9.3 Authentication and Authorization")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "The project uses Clerk for modern authentication and session management.",
            "JWT tokens are verified at protected backend endpoints using middleware before business logic is executed.",
            "Role-based controls are enforced to prevent unauthorized access to educator/admin-only capabilities.",
            "This security model improves trust, data safety, and correctness of user actions across modules.",
        ],
    )

    p = doc.add_paragraph("9.4 Payment and Enrollment Workflow")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "Paid course purchase is integrated with Stripe checkout and webhook verification.",
            "Free course enrollment is handled with a direct server flow while preserving purchase records for audit consistency.",
            "Webhook confirmation ensures enrollment is activated only after successful payment completion in Stripe.",
            "The complete flow minimizes transaction failures and prevents incorrect access provisioning.",
        ],
    )

    p = doc.add_paragraph("9.5 Real-Time Sprint and Leaderboard")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "Socket.IO channels provide real-time events for weekly sprint participation and leaderboard updates.",
            "Students can participate in sprint activities and view rank progression dynamically without full-page refresh.",
            "This module improves engagement through gamified competition and transparent performance feedback.",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("10. Testing, Validation, and Quality Assurance")
    p.style = h1

    p = doc.add_paragraph("10.1 Testing Strategy")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "A multi-layer testing strategy was followed including functional testing, API testing, integration testing, and user acceptance testing.",
            "Key flows tested repeatedly included login/signup, course enrollment, payment confirmation, test attempts, leaderboard updates, and dashboard analytics.",
        ],
    )

    p = doc.add_paragraph("10.2 Sample Functional Test Cases")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "Test Case 1: Verify student login and protected route access.",
            "Test Case 2: Verify paid course purchase and post-webhook enrollment.",
            "Test Case 3: Verify free course enrollment and immediate player access.",
            "Test Case 4: Verify practice test scoring and level progression logic.",
            "Test Case 5: Verify sprint leaderboard updates in real time for active users.",
            "Test Case 6: Verify educator can add, update, publish, and unpublish courses.",
            "Test Case 7: Verify admin analytics pages render expected aggregate values.",
        ],
    )

    p = doc.add_paragraph("10.3 Defect Handling")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "Detected issues were logged and reproduced with clear steps, expected behavior, and actual behavior.",
            "Fixes were validated using regression checks to ensure no side effects on previously stable modules.",
            "Common issues resolved included token header compatibility, route protection edge cases, and payment redirect state handling.",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("11. Results and Discussion")
    p.style = h1

    add_paragraphs(
        doc,
        [
            "The final system successfully demonstrates integrated learning and placement preparation in one platform.",
            "Compared to fragmented approaches, CSEdge improves workflow continuity for students by offering course consumption, practice, and tracking from a single dashboard.",
            "Educators gain improved control and visibility through structured course management and analytics support.",
            "Real-time sprint and leaderboard functionalities increase engagement, while test modules support measurable skill growth.",
            "The architecture remains extensible for future AI-based recommendations and personalized learning paths.",
        ],
    )

    p = doc.add_paragraph("11.1 Screens and Evidence Placeholders")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "[Insert Screenshot: Home Page]",
            "[Insert Screenshot: Student Dashboard]",
            "[Insert Screenshot: Course Details and Player]",
            "[Insert Screenshot: Practice Test Module]",
            "[Insert Screenshot: Arena / Weekly Sprint Leaderboard]",
            "[Insert Screenshot: Educator Course Management]",
            "[Insert Screenshot: Admin Analytics Dashboard]",
            "[Insert Screenshot: Payment Flow Success Page]",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("12. Challenges Faced and Solutions")
    p.style = h1
    add_paragraphs(
        doc,
        [
            "Challenge 1: Managing role-specific routing without exposing restricted pages.",
            "Solution: Implemented protected route wrappers and backend role middleware checks.",
            "Challenge 2: Stable authorization across cross-origin requests.",
            "Solution: Hardened CORS and explicit authorization header handling.",
            "Challenge 3: Reliable payment-to-enrollment synchronization.",
            "Solution: Implemented webhook-driven enrollment updates and status tracking.",
            "Challenge 4: Real-time consistency for leaderboard data.",
            "Solution: Designed socket events and server-side leaderboard recalculation hooks.",
            "Challenge 5: Maintaining user experience across multiple modules.",
            "Solution: Introduced reusable components, clear feedback toasts, and loading states.",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("13. Skills Gained During Internship")
    p.style = h1
    add_paragraphs(
        doc,
        [
            "Full Stack Development: Built and integrated complete frontend-backend workflows.",
            "API Engineering: Designed and consumed secure REST APIs with role-aware authorization.",
            "Database Design: Modeled and optimized MongoDB schemas using Mongoose.",
            "Authentication and Security: Worked with Clerk JWT and middleware-based access controls.",
            "Payment Systems: Implemented Stripe checkout and webhook verification lifecycle.",
            "Real-Time Systems: Developed socket-based weekly sprint and analytics update flows.",
            "Testing and Debugging: Performed defect reproduction, patching, and regression validation.",
            "Project Management: Followed iterative/agile execution with continuous enhancement cycles.",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("14. Future Scope")
    p.style = h1
    add_paragraphs(
        doc,
        [
            "Introduce AI-based personalized learning recommendations based on student performance.",
            "Add adaptive tests that automatically adjust difficulty as users progress.",
            "Develop mobile applications for Android and iOS for broader accessibility.",
            "Integrate proctoring and anti-cheating measures for formal assessment workflows.",
            "Add mentor-mentee matching and advanced interview simulation scenarios.",
            "Enable institutional analytics exports and department-level benchmarking.",
        ],
    )

    doc.add_page_break()
    p = doc.add_paragraph("15. Appendix")
    p.style = h1

    p = doc.add_paragraph("15.1 Sample API Endpoints")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "GET /api/user/data",
            "GET /api/user/enrolled-courses",
            "POST /api/user/purchase",
            "POST /api/user/enroll-free",
            "POST /api/user/update-course-progress",
            "GET /api/course/all",
            "GET /api/course/:id",
            "POST /api/educator/add-course",
            "PUT /api/educator/update-course/:id",
            "GET /api/test/questions",
            "POST /api/test/attempt",
            "GET /api/test/progress/:type",
            "GET /api/sprint/leaderboard",
        ],
    )

    p = doc.add_paragraph("15.2 Sample Data Models (Summary)")
    p.style = h2
    add_paragraphs(
        doc,
        [
            "User Model: identity, role, enrolled courses, profile metadata.",
            "Course Model: title, price, chapters, lectures, educator reference, ratings.",
            "Purchase Model: payment status, amount, session tracking, user and course references.",
            "CourseProgress Model: per-user completion state for chapter/lecture units.",
            "Sprint Model: challenge participation, score, rank metadata, and timestamps.",
        ],
    )

    doc.save(OUTPUT_PATH)
    print(f"Saved final file: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

