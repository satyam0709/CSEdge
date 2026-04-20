from docx import Document


INPUT_PATH = r"D:\ItalentMatch Thesis.docx"
OUTPUT_PATH = r"D:\CSEdge_Internship_Thesis_Satyam_Kumar.docx"


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def main():
    doc = Document(INPUT_PATH)

    # Global phrase-level replacements to keep formatting but shift topic.
    replacements = {
        "ITalentMatch": "CSEdge",
        "Resume Screening and Job Matching Platform": "Full Stack Learning and Placement Platform",
        "recruitment process": "learning and placement preparation process",
        "recruiters": "educators and administrators",
        "candidate": "student",
        "candidates": "students",
        "job matching": "learning recommendation and practice alignment",
        "resume screening": "course and skill progress evaluation",
        "Main Project": "Internship Project",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            replace_paragraph_text(paragraph, updated)

    # Strongly targeted section rewrites (keeps template order/styles).
    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        if not raw:
            continue

        if "Name- " in raw and "Jaydeep Singh Rathore" in raw:
            replace_paragraph_text(paragraph, "Name - Satyam Kumar")
        elif "Enrollment No" in raw:
            replace_paragraph_text(paragraph, "Enrollment No - 2203031050591")
        elif raw.startswith("Internal Guide"):
            replace_paragraph_text(paragraph, "Internal Guide - [Your Faculty Guide Name]")
        elif raw.startswith("Industry Guide"):
            replace_paragraph_text(paragraph, "Industry Guide - [Your Internship Mentor Name]")
        elif raw.startswith("Industry Name"):
            replace_paragraph_text(paragraph, "Industry Name - [Your Internship Company Name], [City], [State]")
        elif raw == "(Resume Screening and Job Matching Platform)":
            replace_paragraph_text(paragraph, "(Full Stack Learning and Placement Platform)")
        elif raw.startswith("We present this document as a comprehensive report"):
            replace_paragraph_text(
                paragraph,
                "We present this document as a comprehensive report of our internship work and progress in developing CSEdge, a Full Stack Learning and Placement Platform. "
                "The system integrates course learning, practice tests, real-time sprint competitions, interview preparation, and educator/admin analytics in a unified web application."
            )
        elif raw.startswith("The system aims to simplify and modernize"):
            replace_paragraph_text(
                paragraph,
                "The system aims to simplify and modernize student preparation by combining learning management, skills practice, and placement-focused tools on one platform. "
                "It bridges the gap between students, educators, and administrators through role-based workflows and data-driven insights."
            )
        elif raw.startswith("This project enhanced our understanding"):
            replace_paragraph_text(
                paragraph,
                "This internship project enhanced my understanding of end-to-end full stack development using React, Node.js, Express, MongoDB, Socket.IO, Clerk authentication, Stripe payments, and Cloudinary integration."
            )
        elif raw.startswith("Currently, most campus recruitment processes are manual and distributed"):
            replace_paragraph_text(
                paragraph,
                "Currently, student learning, coding practice, interview preparation, and progress tracking are fragmented across multiple tools and websites."
            )
        elif raw.startswith("Students create resumes in unstandardized formats."):
            replace_paragraph_text(
                paragraph,
                "Students switch between separate portals for courses, tests, and interview preparation, creating inconsistency and low learning continuity."
            )
        elif raw.startswith("Resume evaluation is often biased"):
            replace_paragraph_text(
                paragraph,
                "Performance tracking is often inconsistent, manual, and not connected to a structured improvement path."
            )
        elif raw.startswith("There is a strong need for an Recruitment platform"):
            replace_paragraph_text(
                paragraph,
                "There is a strong need for an integrated LMS and placement platform that supports guided learning, assessments, and measurable progress."
            )
        elif raw.startswith("Reduce recruiter workload."):
            replace_paragraph_text(paragraph, "Reduce manual effort for educators and administrators through automation and dashboards.")
        elif raw.startswith("Provide students with professional resume templates"):
            replace_paragraph_text(paragraph, "Provide students with structured learning paths, practice modules, and career-readiness tools.")
        elif raw.startswith("Ensure faster and fairer candidate shortlisting."):
            replace_paragraph_text(paragraph, "Ensure faster and fairer student evaluation through standardized tests and analytics.")
        elif raw.startswith("Improve placement efficiency in colleges/universities."):
            replace_paragraph_text(paragraph, "Improve placement readiness and outcomes through continuous practice and feedback loops.")
        elif raw.startswith("Help in resume screening using Logic-driven algorithms."):
            replace_paragraph_text(paragraph, "Build a complete learning and placement ecosystem with course, practice, and sprint modules.")
        elif raw.startswith("Provide intelligent job matching for students based on skills and experience."):
            replace_paragraph_text(paragraph, "Provide topic-wise and level-wise practice modules for aptitude, DSA, development, and SQL.")
        elif raw.startswith("Offer recruiters advanced tools for candidate evaluation and analytics."):
            replace_paragraph_text(paragraph, "Offer educators/admin advanced analytics for engagement, progress, and assessment quality.")
        elif raw.startswith("Enhance student experience with professional resume templates"):
            replace_paragraph_text(paragraph, "Enhance student experience through one dashboard for learning, mock interviews, and resume building.")
        elif raw.startswith("Enable data-driven decisions through analytics dashboards."):
            replace_paragraph_text(paragraph, "Enable data-driven decisions through real-time leaderboards and analytics dashboards.")

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

